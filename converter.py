import fnmatch
import os
import re
import tkinter as tk
import tkinter.ttk as ttk
from pathlib import Path
from tkinter import filedialog as fd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches, Mm, Pt
import openpyxl


# TODO classes to separate files

class RCPDXlsx:
    """ A class to represent an existing RCPD (Register of Processing Operations) Excel document. """

    def __init__(self, folder, filename, read_only):
         """Initialize a RCPDXlsx object."""
        self.folder = folder
        self.filename = filename
        self.path = f'{self.folder}/{self.filename}'
        self.workbook = openpyxl.load_workbook(self.path, data_only=True, read_only=read_only)

    @staticmethod
    def read_row_skipping_odd(worksheet, row):
        """ Return data skipping empty values, which are present due to the merging of columns. """
        return [re.sub('\s+', ' ', cell.value) for i, cell in enumerate(worksheet[row]) if i % 2]

    def extract_data(self, key_row, value_row):
        """ Return data, namely: the filename without extension, the register administrator name,
        the types of regulations and their execution. """
        # do not load cell formulas, values only
        sheet = self.workbook.active
        administrator = sheet['f1'].value
        keys = self.read_row_skipping_odd(sheet, key_row)
        values = self.read_row_skipping_odd(sheet, value_row)
        raw_filename = self.filename.strip('xlsx').strip('.')
        return raw_filename, administrator, keys, values


class NewRCPDDoc:
    """ A class to represent a newly rendered RCPD (Register of Processing Operations) Word document. """

    def __init__(self, folder, raw_filename, administrator, column1, column2, height, width, space,
                 column0_width, column1_width, column2_width):
        """Initialize a NewRCPDDoc object."""
        self.doc = Document()
        self.folder = folder
        self.raw_filename = raw_filename
        self.filename = self.raw_filename + '.docx'
        self.path = f'{self.folder}/{self.filename}'
        self.administrator = administrator
        self.column1 = column1
        self.column2 = column2
        self.height = height
        self.width = width
        self.space = space
        self.column0_width = column0_width
        self.column1_width = column1_width
        self.column2_width = column2_width
        self.table_0_data = [(i, item1, item2) for i, (item1, item2) in enumerate(zip(self.column1, self.column2))]
        self.title = 'Rejestr czynności przetwarzania danych'

    def set_font(self, style_name, font_name, font_size):
        """ Set a name and size of the given style font. """
        style = self.doc.styles[style_name]
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)

    def set_page_size(self):
        """ Set the page size and its equal (but a double top one) margins and spaces."""
        section = self.doc.sections[0]
        section.page_height = Mm(self.height)
        section.page_width = Mm(self.width)
        section.left_margin = Mm(self.space)
        section.right_margin = Mm(self.space)
        section.top_margin = Mm(2 * self.space)
        section.bottom_margin = Mm(self.space)
        section.header_distance = Mm(self.space)
        section.footer_distance = Mm(self.space)

    def set_header(self):
        """ Set the display of a header for all document's pages."""
        header = self.doc.sections[0].header
        header_text = header.paragraphs[0]
        # header_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_text.style = self.doc.styles['Normal']
        header_text.add_run(self.title).bold = True

    def set_subtitle(self):
        """ Set the display of a subtitle. """
        subtitle = self.doc.add_paragraph('Administrator Danych Osobowych  - ')
        # Split the subtitle for a bold span
        subtitle.style = self.doc.styles['Normal']
        subtitle.add_run(self.administrator).bold = True

    def set_table(self):
        """ Add and adjust a table with three columns populated with data. """
        light_grey = 'f2f2f2'
        t = self.doc.add_table(0, 0)
        self.draw_table(t)
        self.populate_table(t, self.table_0_data)
        self.style_table(t, 'Normal')
        self.bold_table_heading(t)
        self.shade_cells(t.columns[0].cells, light_grey)
        self.shade_cells(t.columns[1].cells, light_grey)

    def draw_table(self, table):
        """ Draw a three-column table centered on the page. """
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.add_column(Inches(self.column0_width))
        table.add_column(Inches(self.column1_width))
        table.add_column(Inches(self.column2_width))

    @staticmethod
    def populate_table(table, data):
        """ Populate a three-column table with data. """
        for item0, item1, item2 in data:
            row = table.add_row().cells
            # Convert to string as required datatype
            row[0].text = f'{item0 + 1}.'
            row[1].text = item1
            row[2].text = item2

    def style_table(self, table, style_name):
        """  Modify the style of a table."""
        table.style = 'Table Grid'
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = self.doc.styles[style_name]

    @staticmethod
    def bold_table_heading(table):
        """  Embolden the font of a table."""
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    @staticmethod
    def shade_cells(cells, colour):
        """ Shade given table cells with a given colour."""
        for c in cells:
            tc_v_align = OxmlElement('w:shd')
            tc_v_align.set(qn('w:fill'), colour)
            c._tc.get_or_add_tcPr().append(tc_v_align)

    def modify(self):
        """ Modify the document. """
        self.set_font('Normal', 'Times New Roman', 12)
        self.set_page_size()
        self.set_header()
        self.set_subtitle()
        self.set_table()

    def save(self):
        """ Save the document. """
        # self.doc.save(Path.cwd() / self.folder / filename)
        # self.filename = self.raw_filename + '.docx'
        # self.path = f'{self.folder}/{self.filename}'
        self.doc.save(self.path)


class Application(tk.Frame):
    """ A class to represent GUI application. """
    def __init__(self, master=None):
         """Initialize an Application object."""
        super().__init__(master)
        self.master = master
        # TODO do parametrów, bez hardkodowania
        self.excel_folder = 'excel'
        self.word_folder = 'word'
        self.excel_path = None
        self.word_path = None
        self.grid()
        # self.set_style('W.TLabel', foreground='green')
        # self.set_style('TButton', font=('calibri', 20, 'bold', 'underline'), foreground='red')
        self.create_widgets()


    def select_excel_path(self):
        """ Set the default Excel document path of the application. """
        self.excel_path = fd.askdirectory(initialdir=Path.cwd() / self.excel_folder)

    def select_word_path(self):
         """ Set the default Word document path of the application. """
        self.word_path = fd.askdirectory(initialdir=Path.cwd() / self.word_folder)


    def create_widgets(self):
         """ Create the application widgets. """
        # Create labels
        self.instr_lbl = tk.Label(self,
                                  text='\nBy zmienić któryś z domyślnych katalogów, kliknij odpowiedni przycisk.\n')
        self.instr_lbl.grid(row=0, column=0, columnspan=4, sticky='W')
        self.info_lbl = tk.Label(self, text='')
        self.info_lbl.grid(row=3, column=0, columnspan=4, rowspan=5, sticky='W')
        self.end_lbl = tk.Label(self, text='')
        self.end_lbl.grid(row=8, column=0, columnspan=4, sticky='W')

        # Create buttons
        self.excel_bttn = tk.Button(self, text='Katalog z plikami Excel',
                                    command=self.select_excel_path)
        self.excel_bttn.grid(row=1, column=0, columnspan=2, sticky='W')
        self.word_bttn = tk.Button(self, text='Katalog z plikami Word',
                                   command=self.select_word_path)
        self.word_bttn.grid(row=1, column=2, columnspan=2, sticky='W')
        self.quit_bttn = tk.Button(self, text='Zamknij', fg='red', command=self.master.destroy)
        self.quit_bttn.grid(row=2, column=0, columnspan=2, sticky='W')
        self.submit_buttn = tk.Button(self, text='Konwertuj', fg='green', command=self.convert)
        self.submit_buttn.grid(row=2, column=2, columnspan=2, sticky='W')

    def get_input_data(self):
        """ Set the input location path of the application. """
        if not self.excel_path:
            self.excel_path = Path.cwd() / self.excel_folder
        else:
            if os.path.isdir(self.excel_path):
                self.excel_folder = self.excel_path
            else:
                self.info_lbl.config(
                    text='\nProszę wybrać istniejący katalog z arkuszami Excel.')
        excel_files = fnmatch.filter(os.listdir(self.excel_path), '*xlsx')
        if not excel_files:
            self.info_lbl.config(text='\nKatalog wyjściowy nie zawiera plików z rozszerzeniem "xlsx".')
        return excel_files

    def set_output_location(self):
        """ Set the output location path of the application. """
        if not self.word_path:
            self.word_path = Path.cwd() / self.word_folder
        else:
            if os.path.isdir(self.word_path):
                self.word_folder = self.word_path
            else:
                self.info_lbl.config(
                    text='\nProszę wybrać istniejący katalog na skonwertowane pliki.')

    def create_doc(self, raw_filename, administrator, keys, values):
        """ Create a document with input data. """
        doc = NewRCPDDoc(folder=self.word_folder, raw_filename=raw_filename, administrator=administrator,
                         column1=keys,
                         column2=values, height=297, width=210, space=12.7, column0_width=0.42,
                         column1_width=2.10,
                         column2_width=4.68)
        doc.modify()
        doc.save()

    def convert(self):
        """ Convert an Excel document to the Word format. """
        excel_files = self.get_input_data()
        self.set_output_location()
        for item in excel_files:
            xlsx = RCPDXlsx(folder=self.excel_folder, filename=item, read_only=True)
            raw_filename, administrator, keys, values = xlsx.extract_data(key_row=12, value_row=15)
            self.create_doc(raw_filename, administrator, keys, values)
        self.end_lbl.config(text='\nKonwertowanie zakończone.')


def main():
    root = tk.Tk()
    root.geometry('750x100')
    root.title('Konwertor rejestru: Excel do Word')
    app = Application(master=root)
    app.mainloop()


if __name__ == '__main__':
    main()
