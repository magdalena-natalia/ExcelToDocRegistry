import os
import re
from pathlib import Path
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
# from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from docx.shared import Inches, Mm, Pt
import openpyxl


class RCPDXlsx:
    """ A class to represent an existing RCPD (Register of Processing Operations) Excel document. """

    def __init__(self, folder, filename, read_only):
        self.folder = folder
        self.filename = filename
        self.path = Path.cwd() / self.folder / self.filename
        self.workbook = openpyxl.load_workbook(self.path, data_only=True, read_only=read_only)

    @staticmethod
    def read_row_skipping_odd(worksheet, row):
        """ Return data skipping empty values, which are present due to the merging of columns. """
        return [re.sub('\s+', ' ', cell.value) for i, cell in enumerate(worksheet[row]) if i % 2]

    def extract_data(self, key_row, value_row):
        """ Return data, namely: the filename without extension, the register administrator name,
        the types of regulations and their execution. """
        # do not load cell formulas, values only
        sheet = self.workbook.active
        administrator = sheet['f1'].value
        keys = self.read_row_skipping_odd(sheet, key_row)
        values = self.read_row_skipping_odd(sheet, value_row)
        raw_filename = self.filename.strip('xlsx').strip('.')
        return raw_filename, administrator, keys, values


class NewRCPDDoc:
    """ A class to represent a newly rendered RCPD (Register of Processing Operations) Word document. """

    def __init__(self, folder, raw_filename, administrator, column1, column2, height, width, space,
                 column0_width, column1_width, column2_width):
        self.doc = Document()
        self.folder = folder
        self.raw_filename = raw_filename
        self.administrator = administrator
        self.column1 = column1
        self.column2 = column2
        self.height = height
        self.width = width
        self.space = space
        self.column0_width = column0_width
        self.column1_width = column1_width
        self.column2_width = column2_width
        self.table_0_data = [(i, item1, item2) for i, (item1, item2) in enumerate(zip(self.column1, self.column2))]
        self.title = 'Rejestr czynności przetwarzania danych'

    def set_font(self, style_name, font_name, font_size):
        """ Set a name and size of the given style font. """
        style = self.doc.styles[style_name]
        font = style.font
        font.name = font_name
        font.size = Pt(font_size)

    def set_page_size(self):
        """ Set the page size and its equal (but a double top one) margins and spaces."""
        section = self.doc.sections[0]
        section.page_height = Mm(self.height)
        section.page_width = Mm(self.width)
        section.left_margin = Mm(self.space)
        section.right_margin = Mm(self.space)
        section.top_margin = Mm(2 * self.space)
        section.bottom_margin = Mm(self.space)
        section.header_distance = Mm(self.space)
        section.footer_distance = Mm(self.space)

    def set_header(self):
        """ Set the display of a header for all document's pages."""
        header = self.doc.sections[0].header
        header_text = header.paragraphs[0]
        # header_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_text.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_text.style = self.doc.styles['Normal']
        header_text.add_run(self.title).bold = True

    def set_subtitle(self):
        """ Set the display of a subtitle. """
        subtitle = self.doc.add_paragraph('Administrator Danych Osobowych  - ')
        # Split the subtitle for a bold span
        subtitle.style = self.doc.styles['Normal']
        subtitle.add_run(self.administrator).bold = True

    def set_table(self):
        """ Add and adjust a table with three columns populated with data. """
        light_grey = 'f2f2f2'
        t = self.doc.add_table(0, 0)
        self.draw_table(t)
        self.populate_table(t, self.table_0_data)
        self.style_table(t, 'Normal')
        self.bold_table_heading(t)
        self.shade_cells(t.columns[0].cells, light_grey)
        self.shade_cells(t.columns[1].cells, light_grey)

    def draw_table(self, table):
        """ Draw a three-column table centered on the page. """
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.add_column(Inches(self.column0_width))
        table.add_column(Inches(self.column1_width))
        table.add_column(Inches(self.column2_width))

    @staticmethod
    def populate_table(table, data):
        """ Populate a three-column table with data. """
        for item0, item1, item2 in data:
            row = table.add_row().cells
            # Convert to string as required datatype
            row[0].text = f'{item0 + 1}.'
            row[1].text = item1
            row[2].text = item2

    def style_table(self, table, style_name):
        """  Modify the style of a table."""
        table.style = 'Table Grid'
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.style = self.doc.styles[style_name]

    @staticmethod
    def bold_table_heading(table):
        """  Embolden the font of a table."""
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

    @staticmethod
    def shade_cells(cells, colour):
        """ Shade given table cells with a given colour."""
        for c in cells:
            tc_v_align = OxmlElement('w:shd')
            tc_v_align.set(qn('w:fill'), colour)
            c._tc.get_or_add_tcPr().append(tc_v_align)

    def modify(self):
        """ Modify the document. """
        self.set_font('Normal', 'Times New Roman', 12)
        self.set_page_size()
        self.set_header()
        self.set_subtitle()
        self.set_table()

    def save(self):
        """ Save the document. """
        filename = self.raw_filename + '.docx'
        self.doc.save(Path.cwd() / self.folder / filename)


def main():
    xlsx_folder = 'excel'
    word_folder = 'word'
    # return list of files in directory under the path
    xlsx_files = os.listdir(Path.cwd() / xlsx_folder)
    for item in xlsx_files:
        xlsx = RCPDXlsx(folder=xlsx_folder, filename=item, read_only=True)
        raw_filename, administrator, keys, values = xlsx.extract_data(key_row=12, value_row=15)
        doc = NewRCPDDoc(folder=word_folder, raw_filename=raw_filename, administrator=administrator, column1=keys,
                         column2=values, height=297, width=210, space=12.7, column0_width=0.42, column1_width=2.10,
                         column2_width=4.68)
        doc.modify()
        doc.save()


if __name__ == '__main__':
    main()
